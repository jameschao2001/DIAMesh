"""DIARL — Build USER_GUIDE.docx from docs/USER_GUIDE.md.

Pipeline:
1. Extract every ```mermaid block from the markdown source.
2. For each block, render it to a high-resolution PNG via Playwright +
   the locally installed Chrome / Edge (no network beyond the Mermaid CDN
   that the inline HTML loads). Crop the screenshot to the `<svg>` bbox.
3. Rewrite the markdown so each mermaid block becomes a Markdown image
   reference (`![figN](figures/figN.png)`).
4. Convert the rewritten markdown to HTML with python-markdown
   (extensions: tables, fenced_code, sane_lists, md_in_html).
5. Parse the HTML with BeautifulSoup and emit a python-docx Document
   that mirrors the structure (headings, paragraphs, tables, lists,
   inline runs, code blocks, images).
6. Save to docs/USER_GUIDE.docx.

Author: James Chao, Homi (AI Agent)
Version: 0.1.0
Date: 2026-04-29
"""

from __future__ import annotations

import os
import re
import shutil
import sys
import time
from pathlib import Path

import markdown
from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor, Inches


REPO = Path(__file__).resolve().parent.parent
DOCS = REPO / "docs"
BUILD = REPO / "build"
FIG_DIR = BUILD / "figures"


MERMAID_BATCH_PAGE = """<!DOCTYPE html>
<html><head><meta charset="utf-8">
<script type="module">
  import mermaid from 'https://cdn.jsdelivr.net/npm/mermaid@10/dist/mermaid.esm.min.mjs';
  mermaid.initialize({{ startOnLoad: false, theme: 'default', flowchart: {{ htmlLabels: true, curve: 'basis' }} }});
  (async () => {{
    await mermaid.run({{ querySelector: '.mermaid' }});
    window.__mermaidDone = true;
  }})();
</script>
<style>
  body {{ margin: 0; padding: 24px; background: #ffffff;
          font-family: "Microsoft JhengHei", "Source Han Sans TC", sans-serif; }}
  .mermaid {{ display: block; margin: 32px 0; padding: 16px; background: #ffffff; }}
</style>
</head><body>
{body}
</body></html>
"""


def _find_chrome() -> str | None:
    candidates = [
        os.environ.get("CHROME"),
        shutil.which("chrome"),
        shutil.which("msedge"),
        r"C:\Program Files\Google\Chrome\Application\chrome.exe",
        r"C:\Program Files (x86)\Microsoft\Edge\Application\msedge.exe",
    ]
    for c in candidates:
        if c and Path(c).exists():
            return c
    return None


def render_mermaid_blocks(md_text: str) -> tuple[str, list[Path]]:
    """Render every mermaid block to a PNG and rewrite the markdown.

    Returns the rewritten markdown and the list of rendered PNG paths.
    """
    pattern = re.compile(r"```mermaid\n(.*?)```", re.DOTALL)
    blocks = pattern.findall(md_text)
    if not blocks:
        return md_text, []

    FIG_DIR.mkdir(parents=True, exist_ok=True)

    chrome = _find_chrome()
    if chrome is None:
        raise RuntimeError("No Chromium-based browser found for mermaid rendering")

    # Reuse the HTML produced by build_user_guide_pdf.py — mermaid render
    # is already proven to work there.
    pdf_html = BUILD / "USER_GUIDE.html"
    if not pdf_html.exists():
        print("[render] build/USER_GUIDE.html missing — running build_user_guide_pdf first")
        import subprocess
        subprocess.check_call([sys.executable, "scripts/build_user_guide_pdf.py"])
        if not pdf_html.exists():
            raise RuntimeError("build_user_guide_pdf failed to produce USER_GUIDE.html")

    from playwright.sync_api import sync_playwright

    paths: list[Path] = []
    with sync_playwright() as p:
        browser = p.chromium.launch(
            executable_path=chrome,
            args=["--no-sandbox", "--disable-dev-shm-usage"],
        )
        context = browser.new_context(
            viewport={"width": 1600, "height": 2400}, device_scale_factor=2
        )
        page = context.new_page()
        page.on("console", lambda msg: print(f"[console] {msg.type}: {msg.text[:200]}"))
        page.goto(pdf_html.as_uri(), wait_until="networkidle")
        page.wait_for_function("() => window.__mermaidDone === true", timeout=180000)
        n_present = page.locator(".mermaid").count()
        print(f"[render] mermaid count in HTML: {n_present}, expected: {len(blocks)}")
        for i in range(len(blocks)):
            svg = page.locator(".mermaid").nth(i).locator("svg")
            png_path = FIG_DIR / f"fig{i:02d}.png"
            svg.screenshot(path=str(png_path), omit_background=False)
            paths.append(png_path)
            print(f"[fig {i:02d}] {png_path.name}")
        context.close()
        browser.close()

    def repl(m: re.Match[str]) -> str:
        idx = repl.counter
        repl.counter += 1
        rel = paths[idx].relative_to(REPO).as_posix()
        return f"\n\n![figure {idx + 1}]({rel})\n\n"

    repl.counter = 0
    new_md = pattern.sub(repl, md_text)
    return new_md, paths


# ----- HTML → docx conversion -----

INLINE_TAGS = {"strong", "b", "em", "i", "code", "a", "br"}


def _set_cell_borders(cell):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:color"), "BFBFBF")
        tcBorders.append(b)
    tcPr.append(tcBorders)


def _add_runs(paragraph, node):
    """Walk inline children of `node`, appending styled runs to paragraph."""
    for child in node.children if isinstance(node, Tag) else [node]:
        if isinstance(child, NavigableString):
            text = str(child)
            if text:
                paragraph.add_run(text)
        elif isinstance(child, Tag):
            name = child.name
            if name in {"strong", "b"}:
                run = paragraph.add_run(child.get_text())
                run.bold = True
            elif name in {"em", "i"}:
                run = paragraph.add_run(child.get_text())
                run.italic = True
            elif name == "code":
                run = paragraph.add_run(child.get_text())
                run.font.name = "Consolas"
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)
            elif name == "a":
                # python-docx has no built-in hyperlink helper; render as
                # underlined run with link text. URL kept in text for
                # reference.
                href = child.get("href", "")
                run = paragraph.add_run(child.get_text())
                run.underline = True
                run.font.color.rgb = RGBColor(0x0D, 0x6E, 0xFD)
                if href and href != child.get_text():
                    paragraph.add_run(f" ({href})")
            elif name == "br":
                paragraph.add_run().add_break()
            else:
                # Fall through: walk descendants
                _add_runs(paragraph, child)


def _add_code_block(doc: Document, text: str) -> None:
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Cm(0.5)
    para.paragraph_format.space_before = Pt(4)
    para.paragraph_format.space_after = Pt(4)
    run = para.add_run(text.rstrip("\n"))
    run.font.name = "Consolas"
    run.font.size = Pt(9.5)
    # Light gray background via shading element
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F6F8FA")
    pPr.append(shd)


def _add_list(doc: Document, list_tag: Tag, ordered: bool, level: int = 0) -> None:
    style = "List Number" if ordered else "List Bullet"
    for li in list_tag.find_all("li", recursive=False):
        para = doc.add_paragraph(style=style)
        para.paragraph_format.left_indent = Cm(0.6 * (level + 1))
        # Pull only the text content, but skip nested lists for now
        # (handle them recursively after).
        nested_lists = li.find_all(["ul", "ol"], recursive=False)
        for nl in nested_lists:
            nl.extract()
        _add_runs(para, li)
        # Recurse for nested lists
        for nl in nested_lists:
            _add_list(doc, nl, nl.name == "ol", level + 1)


def _add_table(doc: Document, table_tag: Tag) -> None:
    rows = table_tag.find_all("tr")
    if not rows:
        return
    headers = rows[0].find_all(["th", "td"])
    n_cols = len(headers)
    table = doc.add_table(rows=len(rows), cols=n_cols)
    table.style = "Table Grid"
    table.alignment = WD_ALIGN_PARAGRAPH.LEFT

    for r_idx, row in enumerate(rows):
        cells = row.find_all(["th", "td"])
        for c_idx, cell_tag in enumerate(cells):
            if c_idx >= n_cols:
                break
            cell = table.cell(r_idx, c_idx)
            cell.text = ""  # clear default empty paragraph
            para = cell.paragraphs[0]
            _add_runs(para, cell_tag)
            _set_cell_borders(cell)
            if r_idx == 0:
                # Header row — bold + light fill
                for run in para.runs:
                    run.bold = True
                from docx.oxml import OxmlElement
                from docx.oxml.ns import qn
                tcPr = cell._tc.get_or_add_tcPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:color"), "auto")
                shd.set(qn("w:fill"), "F6F8FA")
                tcPr.append(shd)


def _add_image(doc: Document, src: str, alt: str | None = None) -> None:
    img_path = (REPO / src).resolve()
    if not img_path.exists():
        # Try relative to docs/
        alt_path = (DOCS / src).resolve()
        if alt_path.exists():
            img_path = alt_path
    if not img_path.exists():
        doc.add_paragraph(f"[missing figure: {src}]")
        return
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run()
    # Cap width to roughly fit A4 with 2cm margins (~16cm content width)
    run.add_picture(str(img_path), width=Cm(15))
    if alt:
        cap = doc.add_paragraph(alt)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in cap.runs:
            r.italic = True
            r.font.size = Pt(9)
            r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)


def _walk(doc: Document, node: Tag) -> None:
    for child in node.children:
        if isinstance(child, NavigableString):
            text = str(child).strip()
            if text:
                doc.add_paragraph(text)
            continue
        if not isinstance(child, Tag):
            continue
        name = child.name
        if name == "h1":
            doc.add_heading(child.get_text(), level=0)
        elif name == "h2":
            doc.add_heading(child.get_text(), level=1)
        elif name == "h3":
            doc.add_heading(child.get_text(), level=2)
        elif name == "h4":
            doc.add_heading(child.get_text(), level=3)
        elif name == "p":
            # Image-only paragraphs: emit picture block.
            imgs = child.find_all("img", recursive=False)
            if imgs and len(child.get_text(strip=True)) == 0:
                for img in imgs:
                    _add_image(doc, img.get("src", ""), img.get("alt"))
                continue
            para = doc.add_paragraph()
            _add_runs(para, child)
        elif name == "ul":
            _add_list(doc, child, ordered=False)
        elif name == "ol":
            _add_list(doc, child, ordered=True)
        elif name == "blockquote":
            for sub in child.children:
                if isinstance(sub, Tag) and sub.name == "p":
                    para = doc.add_paragraph()
                    para.paragraph_format.left_indent = Cm(0.6)
                    _add_runs(para, sub)
                    for r in para.runs:
                        r.italic = True
                        r.font.color.rgb = RGBColor(0x5F, 0x45, 0x00)
        elif name == "pre":
            code_tag = child.find("code")
            text = (code_tag or child).get_text()
            _add_code_block(doc, text)
        elif name == "table":
            _add_table(doc, child)
        elif name == "hr":
            doc.add_paragraph("─" * 40).alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif name in {"div", "section", "article"}:
            _walk(doc, child)


def md_to_docx(md_text: str, out_path: Path) -> None:
    html = markdown.markdown(
        md_text,
        extensions=["tables", "fenced_code", "sane_lists", "md_in_html"],
        output_format="html5",
    )
    soup = BeautifulSoup(f"<root>{html}</root>", "html.parser")
    root = soup.root

    doc = Document()
    # Page setup: A4 portrait, 2cm margins
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    # Default font setup for CJK
    style = doc.styles["Normal"]
    style.font.name = "Microsoft JhengHei"
    style.font.size = Pt(11)
    from docx.oxml.ns import qn
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")

    _walk(doc, root)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    print(f"[docx] wrote {out_path.relative_to(REPO)} ({out_path.stat().st_size / 1024:.1f} KB)")


def main() -> int:
    src = DOCS / "USER_GUIDE.md"
    if not src.exists():
        print(f"ERROR: {src} not found", file=sys.stderr)
        return 1

    BUILD.mkdir(exist_ok=True)
    md_text = src.read_text(encoding="utf-8")

    print("[step 1/2] rendering mermaid blocks → PNG ...")
    t0 = time.time()
    md_with_imgs, _ = render_mermaid_blocks(md_text)
    print(f"  done in {time.time() - t0:.1f}s")

    print("[step 2/2] markdown → docx ...")
    t0 = time.time()
    out_path = DOCS / "USER_GUIDE.docx"
    md_to_docx(md_with_imgs, out_path)
    print(f"  done in {time.time() - t0:.1f}s")
    return 0


if __name__ == "__main__":
    sys.exit(main())
